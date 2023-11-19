from PyPDF2 import PdfReader
from docx import Document
from gtts import gTTS
from IPython.display import Audio
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
from peft import PeftModel, PeftConfig
from keybert import KeyBERT
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import CountVectorizer
import numpy as np



# ------------------------------------------------------------Extracting PDF & Docx File-------------------------------------------------------------
def text_extractor(file):
  if file.type=='application/vnd.openxmlformats-officedocument.wordprocessingml.document':
    extracted_docx=[]
    document = Document(file)
    paragraph=document.paragraphs
    for p in document.paragraphs:
      extracted_docx.append(p.text)
    wholeText = ' ' 
    paragraph = wholeText.join(extracted_docx)
    return paragraph
  elif file.type=='application/pdf':
    Extracted_page_list=[]
    text,textN='',''
    reader=PdfReader(file)
    for i in range(len(reader.pages)):
      pageObj = reader.pages[i]
      text=pageObj.extract_text()
      textN=text.replace('\n',' ')
      Extracted_page_list.append(textN)

    wholeText = ' ' 
    result = wholeText.join(Extracted_page_list)
    return result
  else:
    return None


#------------------------------------------------------------------------Fine Tuned BART Model-----------------------------------------------
def fine_tuned_bart(text,max_token):
  config = PeftConfig.from_pretrained("cherryberry01/Re_Sum")
  model = AutoModelForSeq2SeqLM.from_pretrained("facebook/bart-large-cnn")
  model = PeftModel.from_pretrained(model, "cherryberry01/Re_Sum")
  tokenizer=AutoTokenizer.from_pretrained("facebook/bart-large-cnn")

  summ=pipeline('summarization',model=model,tokenizer=tokenizer,truncation=True)
  summary = summ(text,min_length = 10, max_length = max_token)
  return summary[0]['summary_text']

# ----------------------------------------------------------------- Pegasus Model ------------------------------------------------------------
def sum_pegasus(text,min_length):
  
  tokenizer = AutoTokenizer.from_pretrained("google/pegasus-xsum")
  model = AutoModelForSeq2SeqLM.from_pretrained("google/pegasus-xsum")
  text=text
  summarizer = pipeline("summarization", model=model, tokenizer=tokenizer,truncation=True)
  summary=summarizer(text, min_length=min_length, max_length = 10000)
  return summary[0]['summary_text']


#--------------------------------------------------------------------Facebook Large CNN Model------------------------------------------------
def FB_BERT_LARGE_CNN(text,min_length):

  tokenizer = AutoTokenizer.from_pretrained("facebook/bart-large-cnn")
  model = AutoModelForSeq2SeqLM.from_pretrained("facebook/bart-large-cnn")
  text = text
  summarizer = pipeline("summarization", model='facebook/bart-large-cnn', tokenizer=tokenizer,truncation = True)
  summary=summarizer(text, min_length=min_length, max_length = 10000)
  return summary[0]['summary_text']

# ------------------------------------------------------------------ For Translating the Text-------------------------------------------------
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline

def translate(text,target,max_length):
  model = AutoModelForSeq2SeqLM.from_pretrained('facebook/nllb-200-distilled-600M')
  tokenizer = AutoTokenizer.from_pretrained('facebook/nllb-200-distilled-600M')

  translator=pipeline('translation', model=model, tokenizer=tokenizer,src_lang='eng_Latn', tgt_lang=language[target],max_length=max_length)
  translated_text=translator(text)
  return translated_text[0]['translation_text']


language = {'Acehnese (Arabic script)':	'ace_Arab',
            'Acehnese (Latin script)':	'ace_Latn',
'Mesopotamian Arabic':	'acm_Arab',
'Taizzi-Adeni Arabic':	'acq_Arab',
'Tunisian Arabic':	'aeb_Arab',
'Afrikaans':	'afr_Latn',
'South Levantine Arabic':	'ajp_Arab',
'Akan':	'aka_Latn',
'Amharic'	:'amh_Ethi',
'North Levantine Arabic':	'apc_Arab',
'Modern Standard Arabic':	'arb_Arab',
'Modern Standard Arabic (Romanized)':	'arb_Latn',
'Najdi Arabic':	'ars_Arab',
'Moroccan Arabic':	'ary_Arab',
'Egyptian Arabic'	:'arz_Arab',
'Assamese':	'asm_Beng',
'Asturian' :	'ast_Latn',
'Awadhi':	'awa_Deva',
'Central Aymara':	'ayr_Latn',
'South Azerbaijani':	'azb_Arab',
'North Azerbaijani':	'azj_Latn',
'Bashkir':	'bak_Cyrl',
'Bambara':	'bam_Latn',
'Balinese':	'ban_Latn',
'Belarusian':	'bel_Cyrl',
'Bemba':	'bem_Latn',
'Bengali':	'ben_Beng',
'Bhojpuri':	'bho_Deva',
'Banjar (Arabic script)':	'bjn_Arab',
'Banjar (Latin script)':	'bjn_Latn',
'Standard Tibetan':	'bod_Tibt',
'Bosnian':	'bos_Latn',
'Buginese':	'bug_Latn',
'Bulgarian':	'bul_Cyrl',
'Catalan':	'cat_Latn',
'Cebuano':	'ceb_Latn',
'Czech':	'ces_Latn',
'Chokwe':	'cjk_Latn',
'Central Kurdish':	'ckb_Arab',
'Crimean Tatar':	'crh_Latn',
'Welsh':	'cym_Latn',
'Danish':	'dan_Latn',
'German':	'deu_Latn',
'Southwestern Dinka':	'dik_Latn',
'Dyula':	'dyu_Latn',
'Dzongkha'	:'dzo_Tibt',
'Greek':	'ell_Grek',
'English':	'eng_Latn',
'Esperanto':	'epo_Latn',
'Estonian':	'est_Latn',
'Basque':	'eus_Latn',
'Ewe':	'ewe_Latn',
'Faroese':	'fao_Latn',
'Fijian':	'fij_Latn',
'Finnish':	'fin_Latn',
'Fon':	'fon_Latn',
'French':	'fra_Latn',
'Friulian':	'fur_Latn',
'Nigerian Fulfulde':	'fuv_Latn',
'Scottish Gaelic':	'gla_Latn',
'Irish':	'gle_Latn',
'Galician':	'glg_Latn',
'Guarani':	'grn_Latn',
'Gujarati':	'guj_Gujr',
'Haitian Creole':	'hat_Latn',
'Hausa':	'hau_Latn',
'Hebrew':	'heb_Hebr',
'Hindi':	'hin_Deva',
'Chhattisgarhi':	'hne_Deva',
'Croatian':	'hrv_Latn',
'Hungarian':	'hun_Latn',
'Armenian':	'hye_Armn',
'Igbo':	'ibo_Latn',
'Ilocano':	'ilo_Latn',
'Indonesian':	'ind_Latn',
'Icelandic':	'isl_Latn',
'Italian':	'ita_Latn',
'Javanese':	'jav_Latn',
'Japanese':	'jpn_Jpan',
'Kabyle':	'kab_Latn',
'Jingpho':	'kac_Latn',
'Kamba':	'kam_Latn',
'Kannada':	'kan_Knda',
'Kashmiri (Arabic script)':	'kas_Arab',
'Kashmiri (Devanagari script)':	'kas_Deva',
'Georgian':	'kat_Geor',
'Central Kanuri (Arabic script)':	'knc_Arab',
'Central Kanuri (Latin script)':	'knc_Latn',
'Kazakh':	'kaz_Cyrl',
'Kabiyè	': 'kbp_Latn',
'Kabuverdianu':	'kea_Latn',
'Khmer':	'khm_Khmr',
'Kikuyu':	'kik_Latn',
'Kinyarwanda':	'kin_Latn',
'Kyrgyz':	'kir_Cyrl',
'Kimbundu':	'kmb_Latn',
'Northern Kurdish':	'kmr_Latn',
'Kikongo':	'kon_Latn',
'Korean':	'kor_Hang',
'Lao':	'lao_Laoo',
'Ligurian':	'lij_Latn',
'Limburgish':	'lim_Latn',
'Lingala':	'lin_Latn',
'Lithuanian':	'lit_Latn',
'Lombard':	'lmo_Latn',
'Latgalian':	'ltg_Latn',
'Luxembourgish':	'ltz_Latn',
'Luba-Kasai':	'lua_Latn',
'Ganda':	'lug_Latn',
'Luo':	'luo_Latn',
'Mizo':	'lus_Latn',
'Standard Latvian':	'lvs_Latn',
'Magahi':	'mag_Deva',
'Maithili':	'mai_Deva',
'Malayalam':	'mal_Mlym',
'Marathi':	'mar_Deva',
'Minangkabau (Arabic script)':	'min_Arab',
'Minangkabau (Latin script)':	'min_Latn',
'Macedonian':	'mkd_Cyrl',
'Plateau Malagasy':	'plt_Latn',
'Maltese':	'mlt_Latn',
'Meitei (Bengali script)':	'mni_Beng',
'Halh Mongolian	':'khk_Cyrl',
'Mossi':	'mos_Latn',
'Maori':	'mri_Latn',
'Burmese':	'mya_Mymr',
'Dutch':	'nld_Latn',
'Norwegian Nynorsk':	'nno_Latn',
'Norwegian Bokmål':	'nob_Latn',
'Nepali':	'npi_Deva',
'Northern Sotho':	'nso_Latn',
'Nuer':	'nus_Latn',
'Nyanja':	'nya_Latn',
'Occitan':	'oci_Latn',
'West Central Oromo':	'gaz_Latn',
'Odia':	'ory_Orya',
'Pangasinan':	'pag_Latn',
'Eastern Panjabi':	'pan_Guru',
'Papiamento':	'pap_Latn',
'Western Persian':	'pes_Arab',
'Polish	':'pol_Latn',
'Portuguese':	'por_Latn',
'Dari':	'prs_Arab',
'Southern Pashto':	'pbt_Arab',
'Ayacucho Quechua':	'quy_Latn',
'Romanian':	'ron_Latn',
'Rundi':	'run_Latn',
'Russian':	'rus_Cyrl',
'Sango':	'sag_Latn',
'Sanskrit':	'san_Deva',
'Santali':	'sat_Olck',
"Sicilian":	'scn_Latn',
'Shan':	'shn_Mymr',
'Sinhala':	'sin_Sinh',
'Slovak'	:'slk_Latn',
'Slovenian':	'slv_Latn',
'Samoan'	:'smo_Latn',
'Shona':	'sna_Latn',
'Sindhi'	:'snd_Arab',
'Somali'	:'som_Latn',
'Southern Sotho'	:'sot_Latn',
'Spanish'	: 'spa_Latn',
'Tosk Albanian'	:'als_Latn',
'Sardinian'	:'srd_Latn',
  'Tamil' : 'tam_Taml',
  'Telugu' : 'tel_Telu',
  'Kannada' : 'kan_Knda'}
#-----------------------------------------------------------------------Text To Speech---------------------------------------------------
def text_to_speech(text,lang_key):
    tts = gTTS(text=text, lang=tts_language_dict[lang_key])
    tts.save('saved_audio.wav')

tts_language_dict={'Afrikaans':'af',
                  'Arabic':'ar',
                  'Basque':'eu',
                  'Bengali':'bn',
                  'Bulgarian':'bg',
                  'Catalan':'ca',
                  'Chinese':'yue',
                  'Czech':'cs',
                  'Danish':'da',
                  'Dutch':'nl',
                  'English':'en',
                  'Filipino':'fil',
                  'Finnish':'fi',
                  'French':'fr',
                  'Galician':'gl',
                  'German':'de',
                  'Greek':'el',
                  'Gujarati':'gu',
                  'Hebrew':'he',
                  'Hindi':'hi',
                  'Hungarian':'hu',
                  'Icelandic':'is',
                  'Indonesian':'id',
                  'Italian':'it',
                  'Japanese':'ja',
                  'Kannada':'kn',
                  'Korean':'ko',
                  'Latvian':'lv',
                  'Lithuanian':'lt',
                  'Malay':'ms',
                  'Malayalam':'ml',
                  'Mandarin Chinese':'cmn',
                  'Marathi':'mr',
                  'Norwegian':'nb',
                  'Polish':'pl',
                  'Portuguese':'pt',
                  'Punjabi':'pa',
                  'Romanian':'ro',
                  'Russian':'ru',
                  'Serbian':'sr',
                  'Slovak':'sk',
                  'Spanish':'es',
                  'Swedish':'sv',
                  'Tamil':'ta',
                  'Telugu':'te',
                  'Thai':'th',
                  'Turkish':'tr',
                  'Ukrainian':'uk',
                  'Vietnamese':'vi'}



#------------------------------------------------------------KEYWORDS EXTRACTION---------------------------------------------------------------------
def extract_keywords_with_keybert(article_text, top_n=5, diversity=0.5):
    model = KeyBERT(model="distilbert-base-nli-mean-tokens")

    keywords_with_scores = model.extract_keywords(
        article_text,
        top_n=top_n,
        keyphrase_ngram_range=(1, 2),
        stop_words="english",
        diversity=diversity
    )

    # Filter out similar keywords
    filtered_keywords = []
    vectorizer = CountVectorizer().fit([k[0] for k in keywords_with_scores])
    keyword_vectors = vectorizer.transform([k[0] for k in keywords_with_scores])

    for i, keyword in enumerate(keywords_with_scores):
        if i == 0:
            filtered_keywords.append(keyword)
            continue
        sim_scores = cosine_similarity(keyword_vectors[i], keyword_vectors[:i])
        if np.max(sim_scores) < 0.8:  # Threshold for similarity
            filtered_keywords.append(keyword)
    #texts_with_bullet_points = ['• ' + item[0] for item in data]
    texts_only = ['• ' + item[0] for item in filtered_keywords]
    return '\n'.join(texts_only)

